# agents/exporter.py
import os
from typing import Tuple
import markdown2

class ExporterAgent:
    def __init__(self, output_dir: str = "output"):
        self.output_dir = output_dir
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

    def run(self, content: str, title: str) -> Tuple[str, str]:
        """
        Export the content to PDF and DOCX formats.
        
        Args:
            content (str): The formatted article content in Markdown format.
            title (str): The title of the article.
            
        Returns:
            Tuple[str, str]: Paths to the generated PDF and DOCX files.
        """
        # Create DOCX
        docx_path = self._create_docx(content, title)
        
        # Try to create PDF with multiple fallback options
        pdf_path = self._create_pdf_with_fallback(content, title)
        
        return pdf_path, docx_path

    def _create_pdf_with_fallback(self, content: str, title: str) -> str:
        """
        Try multiple methods to create PDF, with fallbacks.
        
        Args:
            content (str): The formatted article content in Markdown format.
            title (str): The title of the article.
            
        Returns:
            str: Path to the generated PDF file (or fallback).
        """
        sanitized_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip().replace(' ', '_')
        pdf_path = os.path.join(self.output_dir, f"{sanitized_title}.pdf")
        
        # Method 1: Try WeasyPrint
        try:
            from weasyprint import HTML
            import re
            html_content = markdown2.markdown(content, extras=['fenced-code-blocks', 'tables'])
            html_content = re.sub(r'<h1.*?>\s*' + re.escape(title) + r'\s*</h1>', '', html_content, flags=re.IGNORECASE)
            # Add enhanced styling with code block borders and better formatting
            styled_html = f"""
            <html>
            <head>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 40px; }}
                    h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
                    h1:first-of-type {{ margin-top: 0; }}
                    h2 {{ color: #34495e; margin-top: 30px; }}
                    h3 {{ color: #7f8c8d; }}
                    pre {{ 
                        background-color: #f8f9fa; 
                        padding: 15px; 
                        border-radius: 8px; 
                        border: 1px solid #dee2e6;
                        overflow-x: auto; 
                        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                        margin: 15px 0;
                    }}
                    code {{ font-family: 'Courier New', monospace; }}
                    p {{ line-height: 1.6; }}                    
                </style>
            </head>
            <body>
                <h1>{title}</h1>
                {html_content}
            </body>
            </html>
            """
            HTML(string=styled_html).write_pdf(pdf_path)
            return pdf_path
            
        except Exception as e:
            print(f"WeasyPrint failed: {e}")
            # Continue to fallback methods
        
        # Method 2: Try ReportLab (if available)
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_LEFT
            
            # Create a simple PDF with ReportLab
            doc = SimpleDocTemplate(pdf_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Add title
            title_para = Paragraph(title, styles['Title'])
            story.append(title_para)
            story.append(Spacer(1, 12))
            
            code_style = ParagraphStyle(
                'Code',
                parent=styles['Normal'],
                fontName='Courier',
                fontSize=10,
                leading=12,
                leftIndent=12,
                rightIndent=12,
                spaceBefore=6,
                spaceAfter=6,
                borderPadding=5,
                borderColor=styles['Normal'].textColor,
                backColor='#f0f0f0'
            )

            # Process content with basic Markdown support
            lines = content.split('\n')
            in_code_block = False
            for line in lines:
                if line.startswith('```'):
                    if in_code_block:
                        # End of a code block
                        para = Paragraph(''.join(code_content).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), code_style)
                        story.append(para)
                        code_content = []
                    in_code_block = not in_code_block
                elif in_code_block:
                    # Accumulate lines within a code block
                    code_content.append(line + '\n')
                elif line.startswith('# ') and title.lower() in line.lower():
                    continue # Skip title, already added
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['h2']))
                elif line.startswith('### '):
                    story.append(Paragraph(line[4:], styles['h3']))
                elif line.strip() and not line.startswith('#'):
                    para = Paragraph(line, styles['Normal'])
                    story.append(para)
                elif not line.strip() and not in_code_block:
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            return pdf_path
            
        except ImportError:
            print("ReportLab not available.")
        except Exception as e:
            print(f"ReportLab failed: {e}")
        
        # Method 3: Fallback to HTML
        import re
        html_path = pdf_path.replace('.pdf', '.html')
        html_content = markdown2.markdown(content, extras=['fenced-code-blocks', 'tables'])
        html_content = re.sub(r'<h1.*?>\s*' + re.escape(title) + r'\s*</h1>', '', html_content, flags=re.IGNORECASE)
        
        # Add enhanced styling with code block borders and better formatting
        styled_html = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; max-width: 800px; margin: 0 auto; padding: 40px; }}
                h1:first-of-type {{ margin-top: 0; }}
                h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
                h2 {{ color: #34495e; margin-top: 30px; }}
                h3 {{ color: #7f8c8d; }}
                pre {{ 
                    background-color: #f8f9fa; 
                    padding: 15px; 
                    border-radius: 8px; 
                    border: 1px solid #dee2e6;
                    overflow-x: auto; 
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                    margin: 15px 0; 
                }}
                code {{ font-family: 'Courier New', monospace; }}
                p {{ line-height: 1.6; }}
            </style>
        </head>
        <body>
            <h1>{title}</h1>
            {html_content}
        </body>
        </html>
        """
        
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(styled_html)
        
        print(f"PDF generation failed. Created HTML file instead: {html_path}")
        return html_path

    def _create_docx(self, markdown_content: str, title: str) -> str:
        """
        Create a DOCX file from Markdown content with improved code snippet handling.
        """
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_COLOR_INDEX
        from docx.shared import Inches
        
        doc = Document()
        # Set document properties
        doc.core_properties.title = title
        doc.core_properties.author = "Multi-Agent Technical Article Generator"
        
        # Add title with specific formatting
        title_para = doc.add_heading(title, 0)
        title_para.style.font.size = Pt(24)
        title_para.style.font.color.rgb = RGBColor(44, 62, 80)  # Dark blue
        
        # Parse markdown content
        lines = markdown_content.split('\n')
        in_code_block = False
        code_language = ""
        code_content = []
        
        for line in lines:
            if line.startswith('```'):
                # Code block start/end
                if in_code_block:
                    # End code block - add formatted code
                    if code_content:
                        # Add a note about the code block
                        if code_language:
                            p = doc.add_paragraph()
                            p.add_run(f"Code example ({code_language}):").italic = True
                        
                        # Add code content with monospace font and better formatting
                        code_para = doc.add_paragraph()
                        code_para.paragraph_format.left_indent = Inches(0.3)
                        code_para.paragraph_format.right_indent = Inches(0.3)
                        code_para.paragraph_format.space_before = Pt(6)
                        code_para.paragraph_format.space_after = Pt(6)
                        
                        run = code_para.add_run()
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                        run.text = '\n'.join(code_content)
                        
                        # Apply light gray background to the paragraph's formatting
                        shd = code_para.paragraph_format.shading
                        # The python-docx library expects a specific format for shading.
                        # This is a way to apply a solid fill color.
                        shd.element.xpath('w:shd')[0].attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill'] = "F0F0F0" # Light Gray
                        # Add spacing after code block
                        spacer = doc.add_paragraph()
                        spacer.paragraph_format.space_before = Pt(12)
                    
                    # Reset code block variables
                    in_code_block = False
                    code_language = ""
                    code_content = []
                else:
                    in_code_block = True
                    code_language = line[3:].strip() if len(line) > 3 else ""
            elif in_code_block:
                code_content.append(line)
            elif line.startswith('# ') and title.lower() in line.lower():
                continue # Skip title, already added as main heading
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.strip():
                doc.add_paragraph(line)
        
        sanitized_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip().replace(' ', '_')
        docx_path = os.path.join(self.output_dir, f"{sanitized_title}.docx")
        doc.save(docx_path)
        return docx_path